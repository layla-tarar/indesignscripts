from __future__ import annotations

import sys
import os
import copy
import re
import tempfile

# Ensure we can import locally installed packages (in .python-packages)
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOCAL_PACKAGES = os.path.join(SCRIPT_DIR, ".python-packages")
if os.path.isdir(LOCAL_PACKAGES) and LOCAL_PACKAGES not in sys.path:
    sys.path.insert(0, LOCAL_PACKAGES)

from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from lxml import etree  # type: ignore


def _make_paragraph_element(text: str, style: str | None = None):
    """Create a bare <w:p> XML element containing the given plain text.
    If style is given, adds <w:pPr><w:pStyle w:val="style"/> so InDesign
    can map the paragraph to the matching paragraph style by name.
    """
    p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _get_description_row(table) -> str | None:
    """Return description text if first row is a merged title row, else None."""
    if not table.rows:
        return None
    first_row_texts = [(cell.text or "").strip() for cell in table.rows[0].cells]
    non_empty = [t for t in first_row_texts if t]
    if non_empty and len(set(non_empty)) == 1:
        return non_empty[0]
    return None


def _extract_description_rows(doc: Document) -> None:
    """
    For each table whose first row is a merged title/description row (all
    non-empty cells contain the same text), insert that text as a plain
    paragraph immediately before the table, then remove the row from the table.
    """
    tbl_tag = qn("w:tbl")
    tr_tag  = qn("w:tr")
    body    = doc.element.body

    tbl_elements  = [el for el in body if el.tag == tbl_tag]
    table_objects = doc.tables  # python-docx objects, same order as XML elements

    for tbl_el, table in zip(tbl_elements, table_objects):
        description = _get_description_row(table)
        if description is None:
            continue
        # Insert description paragraph immediately before the table,
        # tagged with Table_Header so InDesign maps it to that paragraph style.
        tbl_el.addprevious(_make_paragraph_element(description, style="Table_Heading"))
        # Remove the first <w:tr> from the table
        first_tr = tbl_el.find(tr_tag)
        if first_tr is not None:
            tbl_el.remove(first_tr)


def _get_superscript_style_ids(doc: Document) -> frozenset:
    """
    Return a frozenset of character style IDs (w:styleId) whose own rPr defines
    <w:vertAlign w:val="superscript">.  Only checks the style's direct formatting,
    not the basedOn chain.  In practice this catches the built-in Word
    "Footnote Reference" style and any custom superscript character styles.
    """
    result: set[str] = set()
    try:
        styles_el = doc.styles._element
        for style_el in styles_el.findall(qn("w:style")):
            if style_el.get(qn("w:type")) != "character":
                continue
            rpr = style_el.find(qn("w:rPr"))
            if rpr is None:
                continue
            va = rpr.find(qn("w:vertAlign"))
            if va is not None and va.get(qn("w:val")) == "superscript":
                sid = style_el.get(qn("w:styleId"), "")
                if sid:
                    result.add(sid)
    except Exception:
        pass
    return frozenset(result)


def _mark_superscripts(doc: Document) -> None:
    """
    Walk every run in the document (body paragraphs and all table cells).
    - Superscript runs: wrap text as {{text}} and remove the superscript property.
    - Footnote references: replace the <w:footnoteReference> element with {{fn:N}}.
    - Endnote references: replace the <w:endnoteReference> element with {{en:N}}.
    - Character-style superscripts: runs whose character style defines
      <w:vertAlign w:val="superscript"> (e.g. Word's built-in "Footnote Reference"
      style used as a manual superscript marker) are also wrapped as {{text}}.
      Without this, _strip_character_styles would remove the style reference and
      leave the text as plain unstyled body text.
    """
    fn_tag  = qn("w:footnoteReference")
    en_tag  = qn("w:endnoteReference")
    va_tag  = qn("w:vertAlign")
    rpr_tag = qn("w:rPr")

    # Collect character style IDs that define superscript (e.g. "Footnote Reference").
    # Captured by the process_run closure below.
    sup_style_ids = _get_superscript_style_ids(doc)

    def process_run(run):
        r_el = run._r

        # --- Footnote reference ---
        fn_ref = r_el.find(fn_tag)
        if fn_ref is not None:
            fn_id = fn_ref.get(qn("w:id"), "?")
            custom_mark_follows = fn_ref.get(qn("w:customMarkFollows"), "0") == "1"
            # Capture custom mark text BEFORE removing anything.
            # When customMarkFollows="1", the author-typed display text (e.g. "4")
            # sits in a <w:t> in the same run as <w:footnoteReference>.
            existing_texts = [t_el.text for t_el in r_el.findall(qn("w:t")) if t_el.text]
            custom_mark_text = "".join(existing_texts) if existing_texts else None
            r_el.remove(fn_ref)
            rpr = r_el.find(rpr_tag)
            if rpr is not None:
                va = rpr.find(va_tag)
                if va is not None:
                    rpr.remove(va)
            for t_el in r_el.findall(qn("w:t")):
                r_el.remove(t_el)
            t = OxmlElement("w:t")
            if custom_mark_follows and custom_mark_text:
                # Repeat reference (e.g. author reused footnote 4 twice more):
                # emit the display mark as a plain superscript {{4}} so CleanUp.jsx
                # applies Char_Superscript to it, rather than as {{fn:N}} which would
                # become an unmatched marker (the referenced footnote body is empty).
                t.text = f"{{{{{custom_mark_text}}}}}"
            else:
                t.text = f"{{{{fn:{fn_id}}}}}"
            r_el.append(t)
            return

        # --- Endnote reference ---
        en_ref = r_el.find(en_tag)
        if en_ref is not None:
            en_id = en_ref.get(qn("w:id"), "?")
            r_el.remove(en_ref)
            rpr = r_el.find(rpr_tag)
            if rpr is not None:
                va = rpr.find(va_tag)
                if va is not None:
                    rpr.remove(va)
            # Same cleanup for endnote custom marks.
            for t_el in r_el.findall(qn("w:t")):
                r_el.remove(t_el)
            t = OxmlElement("w:t")
            t.text = f"{{{{en:{en_id}}}}}"
            r_el.append(t)
            return

        # --- Direct superscript run ---
        if run.font.superscript:
            run.text = f"{{{{{run.text}}}}}"
            run.font.superscript = False
            return

        # --- Character-style-based superscript ---
        # run.font.superscript only sees direct <w:vertAlign> on the run; it returns
        # None when superscript is inherited from a character style.  Check the
        # applied character style ID against the pre-built set of superscript styles.
        if sup_style_ids and run.text.strip():
            rpr = r_el.find(rpr_tag)
            if rpr is not None:
                r_style = rpr.find(qn("w:rStyle"))
                if r_style is not None and r_style.get(qn("w:val"), "") in sup_style_ids:
                    run.text = f"{{{{{run.text}}}}}"
                    # Leave w:rStyle in place — _strip_character_styles removes it.

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                process_run(run)

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)


def _strip_field_codes(doc: Document) -> None:
    """
    Remove Word field code machinery from all paragraphs, keeping only the
    display text (the portion between 'separate' and 'end' markers).
    """
    fldchar_tag      = qn("w:fldChar")
    r_tag            = qn("w:r")
    fldchartype_attr = qn("w:fldCharType")

    def strip_fields(p_el):
        in_instruction = False
        for child in list(p_el):
            if child.tag != r_tag:
                continue
            fldchar = child.find(fldchar_tag)
            if fldchar is not None:
                ftype = fldchar.get(fldchartype_attr, "")
                if ftype == "begin":
                    in_instruction = True
                    p_el.remove(child)
                elif ftype == "separate":
                    in_instruction = False
                    p_el.remove(child)
                elif ftype == "end":
                    p_el.remove(child)
            elif in_instruction:
                p_el.remove(child)

    for para in doc.paragraphs:
        strip_fields(para._p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    strip_fields(para._p)


def _strip_character_styles(doc: Document) -> None:
    """
    Strip run-level formatting that would override InDesign paragraph styles on placement:
    - Unwrap <w:hyperlink> elements (keeping link text as plain runs)
    - Remove from <w:rPr>: rStyle, rFonts, sz/szCs, color, b/bCs, u, highlight, kern,
      spacing, lang
    Preserves: i/iCs (italics) and vertAlign (superscript — handled separately).
    """
    rpr_tag       = qn("w:rPr")
    hyperlink_tag = qn("w:hyperlink")

    _STRIP_RPR = {
        qn("w:rStyle"),
        qn("w:rFonts"),
        qn("w:sz"),    qn("w:szCs"),
        qn("w:color"),
        qn("w:b"),     qn("w:bCs"),
        qn("w:u"),
        qn("w:highlight"),
        qn("w:kern"),
        qn("w:spacing"),
        qn("w:lang"),
    }

    ppr_tag = qn("w:pPr")

    def strip_run_overrides(paragraphs):
        for para in paragraphs:
            p_el = para._p
            for hl in p_el.findall(hyperlink_tag):
                parent = hl.getparent()
                idx = list(parent).index(hl)
                for child in list(hl):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(hl)

            ppr = p_el.find(ppr_tag)
            if ppr is not None:
                ppr_rpr = ppr.find(rpr_tag)
                if ppr_rpr is not None:
                    ppr.remove(ppr_rpr)

            for run in para.runs:
                rpr = run._r.find(rpr_tag)
                if rpr is not None:
                    for child in list(rpr):
                        if child.tag in _STRIP_RPR:
                            rpr.remove(child)

    strip_run_overrides(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                strip_run_overrides(cell.paragraphs)


def _clean_runs(doc: Document) -> None:
    """
    Text-level cleanup:
    - Strip leading bullet character (U+2022 + space) from paragraph starts.
    - Replace tilde operator (U+223C) with standard tilde (~).
    - Normalize lone "X" in table cells to lowercase "x".
    """
    BULLET = "\u2022"
    TILDE_OP = "\u223c"

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.text.startswith(BULLET + " "):
                    run.text = run.text[2:]
                elif run.text.startswith(BULLET):
                    run.text = run.text[1:]
                break

            for run in para.runs:
                if TILDE_OP in run.text:
                    run.text = run.text.replace(TILDE_OP, "~")

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
                if cell.text.strip() == "X":
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace("X", "x")


_FOOTNOTES_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
_ENDNOTES_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
)
_SKIP_FOOTNOTE_TYPES = {"separator", "continuationSeparator", "continuationNotice"}


def extract_footnotes(source_path: str, output_path: str) -> int:
    """
    Extract all footnotes/endnotes from source_path into a new .docx.
    Returns the number of notes written.
    """
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    doc = Document(source_path)
    out_doc = Document()
    body = out_doc.element.body
    count = 0

    default_p = body.find(qn("w:p"))
    if default_p is not None:
        body.remove(default_p)

    fn_tag    = qn("w:footnote")
    en_tag    = qn("w:endnote")
    id_attr   = qn("w:id")
    type_attr = qn("w:type")
    t_tag     = qn("w:t")
    ref_tags  = {qn("w:footnoteRef"), qn("w:endnoteRef")}
    rpr_tag   = qn("w:rPr")
    keep_rpr  = {qn("w:i"), qn("w:iCs")}

    def _copy_note_paragraphs(note_el):
        for p_el in note_el.findall(qn("w:p")):
            p_copy = copy.deepcopy(p_el)
            for r_el in p_copy.findall(qn("w:r")):
                non_rpr = [c for c in list(r_el) if c.tag != rpr_tag]
                if len(non_rpr) == 1 and non_rpr[0].tag in ref_tags:
                    r_el.getparent().remove(r_el)
            remaining = p_copy.findall(qn("w:r"))
            if remaining:
                t_el = remaining[0].find(t_tag)
                if t_el is not None and t_el.text:
                    stripped = t_el.text.lstrip(" \t\xa0")
                    if stripped:
                        t_el.text = stripped
                    else:
                        remaining[0].getparent().remove(remaining[0])
            for r_el in p_copy.findall(qn("w:r")):
                rpr = r_el.find(rpr_tag)
                if rpr is not None:
                    for child in list(rpr):
                        if child.tag not in keep_rpr:
                            rpr.remove(child)
                    if len(rpr) == 0:
                        r_el.remove(rpr)
            body.append(p_copy)

    def _extract_from_part(rel_uri, note_tag, label_prefix):
        nonlocal count
        try:
            part = doc.part.part_related_by(rel_uri)
            part_xml = etree.fromstring(part.blob)
        except (KeyError, Exception):
            return
        for note in part_xml.findall(note_tag):
            if note.get(type_attr) in _SKIP_FOOTNOTE_TYPES:
                continue
            note_id = note.get(id_attr, "?")
            body.append(_make_paragraph_element(f"{{{{{label_prefix}:{note_id}}}}}"))
            _copy_note_paragraphs(note)
            body.append(_make_paragraph_element(""))
            count += 1

    _extract_from_part(_FOOTNOTES_REL, fn_tag, "fn")
    _extract_from_part(_ENDNOTES_REL,  en_tag, "en")

    if count == 0:
        return 0

    out_doc.save(output_path)
    return count


def export_footnotes_txt(source_path: str, output_path: str) -> int:
    """
    Write a tab-separated .txt mapping footnote/endnote IDs to plain text.
    Returns the number of entries written.
    """
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    doc = Document(source_path)

    t_tag     = qn("w:t")
    rpr_tag   = qn("w:rPr")
    fn_tag    = qn("w:footnote")
    en_tag    = qn("w:endnote")
    id_attr   = qn("w:id")
    type_attr = qn("w:type")
    ref_tags  = {qn("w:footnoteRef"), qn("w:endnoteRef")}

    entries: list[tuple[str, str]] = []

    def _note_plain_text(note_el) -> str:
        paragraphs = []
        for p_el in note_el.findall(qn("w:p")):
            para_text = ""
            for r_el in p_el.findall(qn("w:r")):
                non_rpr = [c for c in list(r_el) if c.tag != rpr_tag]
                if len(non_rpr) == 1 and non_rpr[0].tag in ref_tags:
                    continue
                for t_el in r_el.findall(t_tag):
                    if t_el.text:
                        para_text += t_el.text
            stripped = para_text.lstrip(" \t\xa0")
            if stripped:
                paragraphs.append(stripped)
        return "\r".join(paragraphs)

    def _collect_from_part(rel_uri, note_tag, label_prefix):
        try:
            part = doc.part.part_related_by(rel_uri)
            part_xml = etree.fromstring(part.blob)
        except (KeyError, Exception):
            return
        for note in part_xml.findall(note_tag):
            if note.get(type_attr) in _SKIP_FOOTNOTE_TYPES:
                continue
            note_id = note.get(id_attr, "?")
            text = _note_plain_text(note)
            if text:
                entries.append((f"{label_prefix}:{note_id}", text))

    _collect_from_part(_FOOTNOTES_REL, fn_tag, "fn")
    _collect_from_part(_ENDNOTES_REL,  en_tag, "en")

    if not entries:
        return 0

    with open(output_path, "w", encoding="utf-8") as fh:
        for key, text in entries:
            fh.write(f"{key}\t{text}\n")

    return len(entries)


# ---------------------------------------------------------------------------
# Heuristic paragraph style inference
# ---------------------------------------------------------------------------

# InDesign style names assigned by this function — used to skip already-styled
# paragraphs on subsequent passes.
_ASSIGNED_STYLES = frozenset({
    "Table_Heading", "Table_FootNote", "Head_SubsectionUnnumbered",
    "Head_CropName", "Head_Section",
})

# Word paragraph style IDs that carry no structural meaning — heuristics only
# run on paragraphs with one of these styles (or no style at all).
# Paragraphs with any other Word style (e.g. Heading1, Title) are left alone
# so that CleanUp.jsx can remap them to the correct InDesign style.
_WORD_BODY_STYLE_IDS = frozenset({
    None,
    "Normal", "Normal0",
    "Body Text", "BodyText", "BodyText1", "Body Text1",
    "Default", "DefaultParagraphFont", "Default Paragraph Font",
})

_TABLE_CAPTION_RE  = re.compile(r"^Table\s+\d+[.:]", re.IGNORECASE)
_TRAILING_PAREN_RE = re.compile(r"\s*\([^)]+\)\s*$")
# Table footnote marker patterns (non-superscripted — superscripted markers are
# already caught by the {{...}} check):
#   * or ** at the start of a paragraph  →  e.g. "*Pakistan has also..."
#   single lowercase letter + space      →  e.g. "a Values represent..."
#     (paired with sz < body default to avoid matching normal body sentences)
_TABLE_FOOTNOTE_ASTERISK_RE    = re.compile(r"^\*+\s*\S")
_TABLE_FOOTNOTE_LETTER_MARK_RE = re.compile(r"^[a-z]\s")


def _get_para_style_val(para) -> str | None:
    """Return the w:pStyle val for this paragraph, or None if not set."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _set_para_style_val(para, style_name: str) -> None:
    """Stamp a w:pStyle value onto an existing paragraph element."""
    p = para._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_name)


def _first_run_is_bold_or_smallcaps(para) -> bool:
    """Return True if the first non-empty run has bold or small-caps formatting."""
    for run in para.runs:
        if not run.text.strip():
            continue
        rpr = run._r.find(qn("w:rPr"))
        if rpr is None:
            return False
        for tag in (qn("w:b"), qn("w:smallCaps")):
            el = rpr.find(tag)
            if el is not None and el.get(qn("w:val"), "true") not in ("false", "0"):
                return True
        return False
    return False


def _first_run_font_size(para) -> int | None:
    """Return the w:sz half-point value of the first non-empty run, or None if unset."""
    for run in para.runs:
        if not run.text.strip():
            continue
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            sz_el = rpr.find(qn("w:sz"))
            if sz_el is not None:
                val = sz_el.get(qn("w:val"))
                if val is not None:
                    return int(val)
        return None  # first non-empty run has no explicit sz
    return None


def _first_run_is_italic(para) -> bool:
    """Return True if the first non-empty run has italic formatting."""
    for run in para.runs:
        if not run.text.strip():
            continue
        rpr = run._r.find(qn("w:rPr"))
        if rpr is None:
            return False
        el = rpr.find(qn("w:i"))
        if el is not None and el.get(qn("w:val"), "true") not in ("false", "0"):
            return True
        return False
    return False


def _register_assigned_styles(doc: Document) -> None:
    """
    Add a minimal Word paragraph style definition to styles.xml for each
    InDesign style name we assign via _set_para_style_val.

    Without a definition in styles.xml the <w:pStyle w:val="..."> reference
    is unknown to InDesign, which falls back to [Basic Paragraph] and then
    CleanUp.jsx converts those paragraphs to Body_Text.

    The styles are custom (w:customStyle="1"), based on Normal, and carry no
    formatting of their own — InDesign maps them to the matching InDesign
    paragraph style by name on placement.
    """
    try:
        styles_el = doc.styles._element
    except AttributeError:
        return

    existing_ids = {
        el.get(qn("w:styleId"), "")
        for el in styles_el.findall(qn("w:style"))
    }

    for style_name in sorted(_ASSIGNED_STYLES):
        if style_name in existing_ids:
            continue
        w_style = OxmlElement("w:style")
        w_style.set(qn("w:type"), "paragraph")
        w_style.set(qn("w:customStyle"), "1")
        w_style.set(qn("w:styleId"), style_name)
        w_name = OxmlElement("w:name")
        w_name.set(qn("w:val"), style_name)
        w_style.append(w_name)
        w_basedOn = OxmlElement("w:basedOn")
        w_basedOn.set(qn("w:val"), "Normal")
        w_style.append(w_basedOn)
        styles_el.append(w_style)


def _infer_paragraph_styles(doc: Document) -> None:
    """
    Heuristically assign InDesign paragraph styles to body-text paragraphs.

    Must run AFTER _mark_superscripts (needs {{...}} markers in text) and
    BEFORE _strip_character_styles (needs bold/small-caps XML attributes).

    Priority (first match wins; already-assigned paragraphs are skipped):
      1. Table_Header              — "Table N." or "Table N:"
      2. Table_FootNote            — starts with {{ but not a {{fn:}} / {{en:}} ref
                                     (superscripted markers converted by _mark_superscripts)
      3. Table_FootNote            — starts with one or more * (literal asterisk markers)
      4. Table_FootNote            — starts with a lowercase letter + space AND font size
                                     < 24 half-pts (catches "a Values…" / "b Note…" markers
                                     that are not superscripted in the source)
      4b.Table_FootNote            — starts with a lowercase letter + space AND first run
                                     is italic at body size (no explicit sz, but italic
                                     signals annotation rather than body text)
      5. Head_SubsectionUnnumbered — first run is bold or small-caps, not all caps
      6. Head_SubsectionUnnumbered — first run font size > 24 half-pts (> 12pt body default),
                                     not all caps (catches enlarged-font headings with no other
                                     formatting signal)
      7. Head_CropName             — 1–2 all-caps words (+ optional parenthetical), no period
      8. Head_Section              — 3+ all-caps words, no period
    """
    for para in doc.paragraphs:
        style_val = _get_para_style_val(para)

        # Skip paragraphs we already assigned a style to on a previous pass.
        if style_val in _ASSIGNED_STYLES:
            continue

        # Skip paragraphs that already carry a meaningful Word structural style
        # (Heading 1/2/3, Title, List Paragraph, etc.).  CleanUp.jsx remaps
        # those to the correct InDesign styles, so the heuristics below must
        # not overwrite them.
        if style_val not in _WORD_BODY_STYLE_IDS:
            continue

        text = para.text.strip()
        if not text:
            continue

        # 1. Table caption / header
        if _TABLE_CAPTION_RE.match(text):
            _set_para_style_val(para, "Table_Heading")
            continue

        # 2. Table footnote — {{letter/number}} superscript marker (not a fn/en ref)
        if text.startswith("{{") and not (
            text.startswith("{{fn:") or text.startswith("{{en:")
        ):
            _set_para_style_val(para, "Table_FootNote")
            continue

        # 3. Table footnote — literal asterisk marker (* or **)
        if _TABLE_FOOTNOTE_ASTERISK_RE.match(text):
            _set_para_style_val(para, "Table_FootNote")
            continue

        # 4. Table footnote — single lowercase letter + space, sub-body font size
        # e.g. "a Values represent the means..." (marker not superscripted in source)
        sz = _first_run_font_size(para)
        if _TABLE_FOOTNOTE_LETTER_MARK_RE.match(text) and sz is not None and sz < 24:
            _set_para_style_val(para, "Table_FootNote")
            continue

        # 4b. Table footnote — single lowercase letter + space, italic at body size
        # Catches the same "a Note…" pattern when no explicit font size is set but
        # the run is italic (annotation styling at the default 12pt).
        if _TABLE_FOOTNOTE_LETTER_MARK_RE.match(text) and sz is None and _first_run_is_italic(para):
            _set_para_style_val(para, "Table_FootNote")
            continue

        # 5. Bold/small-caps heading (not all caps — those are handled below)
        if not text.isupper() and _first_run_is_bold_or_smallcaps(para):
            _set_para_style_val(para, "Head_SubsectionUnnumbered")
            continue

        # 6. Enlarged-font heading (not all caps).
        # Catches paragraphs where the author simply increased the font size above
        # the 12pt body default (sz=24 half-pts) without applying a Word heading style
        # or bold/small-caps.  All-caps text is excluded here so it falls through to
        # the Head_CropName / Head_Section heuristics below.
        # sz was already fetched above (step 4); call again only if not yet set.
        if sz is None:
            sz = _first_run_font_size(para)
        if sz is not None and sz > 24 and not text.isupper():
            _set_para_style_val(para, "Head_SubsectionUnnumbered")
            continue

        # 7 & 8. All-caps section / crop-name headers
        if text.endswith("."):
            continue

        # Strip trailing parenthetical "(Zea mays)" before word counting
        core = _TRAILING_PAREN_RE.sub("", text).strip()
        if not core or not core.isupper():
            continue

        words = core.split()
        if 1 <= len(words) <= 2:
            _set_para_style_val(para, "Head_CropName")
        elif len(words) >= 3:
            _set_para_style_val(para, "Head_Section")


# ---------------------------------------------------------------------------
# Per-monograph new-term scanner
# ---------------------------------------------------------------------------

# Latin terms already handled by CleanUp.jsx — omit from all scan categories.
_KNOWN_LATIN = frozenset({
    "in vitro", "in vivo", "in situ", "de novo", "ex vivo", "ad libitum",
    "in planta", "in silico", "sensu stricto", "sensu lato",
    "et al", "et al.", "per se", "in utero", "in ovo",
})

# Scientific names (binomial + abbreviated genus) already seen — lowercase.
_KNOWN_SCI_NAMES = frozenset({
    "bacillus thuringiensis", "bacillus thuringiensis var. kurstaki",
    "bacillus thuringiensis subsp. kurstaki",
    "zea mays", "glycine max", "brassica napus", "gossypium hirsutum",
    "oryza sativa", "solanum tuberosum", "arabidopsis thaliana",
    "nicotiana tabacum", "saccharomyces cerevisiae", "escherichia coli",
    "helicoverpa armigera", "helicoverpa zea",
    "spodoptera exigua", "spodoptera frugiperda",
    "manduca sexta", "ostrinia nubilalis",
    "bombyx mori", "plutella xylostella",
    "diatraea saccharalis", "agrotis ipsilon",
    # abbreviated genus forms
    "b. thuringiensis", "e. coli", "z. mays",
})

# Gene names already seen (lowercase for comparison).
_KNOWN_GENE_NAMES = frozenset({
    "cry1ab", "cry1ac", "cry1a", "cry1b", "cry1c", "cry1d", "cry1e", "cry1f",
    "cry2ab", "cry2ae", "cry2a",
    "cry3bb", "cry3a", "cry3b",
    "cry34ab1", "cry35ab1",
    "cry9c",
    "epsps", "bar", "pat", "gat",
    "nptii", "aad", "gus", "hpt",
    "vip3a", "vip3aa", "vip3ab",
})

# Protein names already seen (lowercase for comparison).
_KNOWN_PROTEIN_NAMES = frozenset({
    "cry1ab", "cry1ac", "cry2ab", "cry3bb", "cry1f",
    "cry34ab1", "cry35ab1",
    "epsps", "cp4 epsps", "bar", "pat",
    "vip3a", "vip3aa",
})

# Regulatory transformation event names already seen (lowercase).
_KNOWN_EVENT_NAMES = frozenset({
    "mon810", "mon863", "mon87701", "mon89034", "mon87460",
    "nk603", "ga21",
    "t25", "ll62",
    "bt11", "bt176",
    "mir162", "mir604",
    "das-44406-6", "das-59122-7",
    "syn-e3272-5", "syn-ir102-7",
    "mzir098", "smg-00355-3",
})

# Common abbreviations to suppress in the regulatory abbreviation scan.
_KNOWN_REG_ABBREVS = frozenset({
    # Countries / regions
    "us", "eu", "uk", "un", "ca", "jp", "au", "nz", "br",
    # Regulatory bodies
    "fao", "who", "epa", "fda", "usda", "efsa", "oecd", "codex",
    "aphis", "ec", "echa", "cfia", "fsanz", "bfr", "rivm", "anses",
    # Molecular biology
    "dna", "rna", "mrna", "rrna", "trna", "dsrna", "sirna", "rnai", "cdna",
    "pcr", "qpcr", "elisa", "hplc", "sds",
    # Dose/safety metrics
    "lc50", "ld50", "noec", "noael", "loael", "loec", "bw", "adi", "arfd", "tdi",
    "auc", "gras",
    # Units
    "bp", "kb", "mb", "kda", "da", "ppm", "ppb", "ppt", "mg", "kg", "ng",
    # Project / document
    "ffs", "afsi", "gmo", "lmo", "gm", "ge", "bt",
    # GLP / GMP / regulatory documents
    "glp", "gmp", "gcp", "sop",
    # International agreements
    "ippc", "cbd", "cpb",
    # Other common bio terms
    "cas", "ntp",
    # Combined-form country+agency tokens
    "usepa", "usfda",
    # Molecular biology terms common in safety dossiers
    "utr", "ssu", "lsu", "orf", "snp", "ssr", "indel",
    # In vitro digestion / safety-study abbreviations
    "sgf", "sif", "bsa",
    # Roman numerals (appear in section numbering and table entries)
    "iii", "vii", "viii", "xii", "xiii", "xiv", "xvi",
})

# --- Regex patterns ---

# Binomial: "Bacillus thuringiensis", "Spodoptera frugiperda"
# Require species epithet ≥6 chars to avoid matching common English words.
_BINOMIAL_RE     = re.compile(r'\b([A-Z][a-z]{2,})\s+([a-z]{6,})\b')
# Abbreviated genus: "B. thuringiensis"
_ABBREV_GENUS_RE = re.compile(r'\b([A-Z])\.\s+([a-z]{3,})\b')
# Gene name token from italic text: cry1Ab, epsps, bar, nptII
_GENE_TOKEN_RE   = re.compile(r'\b([a-z]{2,5}\d[A-Za-z0-9]*)\b')
# Cry / Vip protein names in full text.
# Cap alpha suffix at 5 chars so the word-boundary check blocks
# false matches like "Cry1AbProtein" (no \b mid-word).
_CRY_PROTEIN_RE  = re.compile(
    r'\b(Cry\d+[A-Za-z]{1,5}\d*|Vip\d+[A-Za-z]{1,5}\d*)\b'
)
# Regulatory event names (conservative prefixes only)
_EVENT_RE = re.compile(
    r'\b('
    r'MON\d{2,}'                       # MON810, MON87701
    r'|NK\d{2,}'                       # NK603
    r'|GA\d{2,}'                       # GA21
    r'|MIR\d{2,}'                      # MIR162, MIR604
    r'|LL\d{2,}'                       # LL62
    r'|Bt\d{2,}'                       # Bt11, Bt176
    r'|T\d{2}'                         # T25 (2-digit avoids T1/T2 timepoint refs)
    r'|[A-Z]{2,5}-\d{4,}(?:-\d+)+'    # DAS-44406-6
    r'|[A-Z]{2,5}-[A-Z]\d{3,}-\d'     # SYN-E3272-5
    r')\b'
)
# All-caps abbreviations 3–6 chars for regulatory body scan
_REG_ABBREV_RE = re.compile(r'\b([A-Z]{3,6})\b')


def _iter_all_paragraphs(doc: Document):
    """Yield all paragraphs from body paragraphs and table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _collect_italic_spans(doc: Document) -> list[str]:
    """
    Return text spans built by joining consecutive italic runs within each
    paragraph.  Used to detect scientific names and gene names before
    _strip_character_styles removes italic markup.
    """
    i_tag   = qn("w:i")
    rpr_tag = qn("w:rPr")
    spans: list[str] = []

    for para in _iter_all_paragraphs(doc):
        current: list[str] = []
        for run in para.runs:
            rpr    = run._r.find(rpr_tag)
            italic = False
            if rpr is not None:
                i_el = rpr.find(i_tag)
                if i_el is not None and i_el.get(qn("w:val"), "true") not in ("false", "0"):
                    italic = True
            if italic and run.text:
                current.append(run.text)
            else:
                if current:
                    span = "".join(current).strip()
                    if span:
                        spans.append(span)
                    current = []
        if current:
            span = "".join(current).strip()
            if span:
                spans.append(span)

    return spans


def _collect_body_texts(doc: Document) -> list[str]:
    """Return paragraph text strings, skipping all-caps heading paragraphs."""
    return [
        para.text.strip()
        for para in _iter_all_paragraphs(doc)
        if para.text.strip() and not para.text.strip().isupper()
    ]


def scan_new_terms(source_path: str) -> dict[str, list[str]]:
    """
    Scan the original .docx for candidate new terms not in the known lists.
    Must be called on the unmodified source so italic markup is still present.
    Returns: { category_name: sorted list of candidate strings }
    """
    doc = Document(source_path)

    italic_spans  = _collect_italic_spans(doc)
    body_texts    = _collect_body_texts(doc)
    italic_joined = " | ".join(italic_spans)
    body_joined   = "\n".join(body_texts)

    # Full text for protein / event scanning (body + footnotes plain text)
    full_text = body_joined
    try:
        part      = doc.part.part_related_by(_FOOTNOTES_REL)
        part_xml  = etree.fromstring(part.blob)
        fn_tag    = qn("w:footnote")
        t_tag     = qn("w:t")
        type_attr = qn("w:type")
        for note in part_xml.findall(fn_tag):
            if note.get(type_attr) in _SKIP_FOOTNOTE_TYPES:
                continue
            for t_el in note.findall(".//" + t_tag):
                if t_el.text:
                    full_text += "\n" + t_el.text
    except Exception:
        pass

    # 1. Scientific names — from italic spans only (reduces false positives)
    # Additional guard: skip if the first word is a common English non-genus word
    # (catches e.g. "Values represent", "Results indicate" from table footnotes).
    _NON_GENUS = frozenset({
        "values", "results", "table", "figure", "note", "these", "those",
        "when", "which", "where", "data", "mean", "means", "each", "both",
        "percent", "percentage", "numbers", "based", "given", "shown",
        "using", "used", "total", "average",
    })
    new_sci: set[str] = set()
    for m in _BINOMIAL_RE.finditer(italic_joined):
        if m.group(1).lower() in _NON_GENUS:
            continue
        term = f"{m.group(1)} {m.group(2)}"
        low  = term.lower()
        if low not in _KNOWN_SCI_NAMES and low not in _KNOWN_LATIN:
            new_sci.add(term)
    for m in _ABBREV_GENUS_RE.finditer(italic_joined):
        term = f"{m.group(1)}. {m.group(2)}"
        if term.lower() not in _KNOWN_SCI_NAMES:
            new_sci.add(term)

    # 2. Gene names — from italic spans, lowercase pattern
    new_genes: set[str] = set()
    for span in italic_spans:
        for m in _GENE_TOKEN_RE.finditer(span):
            cand = m.group(1)
            if cand.lower() not in _KNOWN_GENE_NAMES and cand.lower() not in _KNOWN_LATIN:
                new_genes.add(cand)

    # 3. Protein names — Cry/Vip pattern from full text
    new_proteins: set[str] = set()
    for m in _CRY_PROTEIN_RE.finditer(full_text):
        cand = m.group(1)
        if cand.lower() not in _KNOWN_PROTEIN_NAMES:
            new_proteins.add(cand)

    # 4. Event names — conservative prefix patterns from full text
    new_events: set[str] = set()
    for m in _EVENT_RE.finditer(full_text):
        cand = m.group(1)
        if cand.lower() not in _KNOWN_EVENT_NAMES:
            new_events.add(cand)

    # 5. Regulatory abbreviations — all-caps 3–6 chars, ≥2 occurrences in
    #    mixed-case body text (skips all-caps heading paragraphs already filtered
    #    by _collect_body_texts)
    abbrev_counts: dict[str, int] = {}
    for m in _REG_ABBREV_RE.finditer(body_joined):
        cand = m.group(1)
        if cand.lower() not in _KNOWN_REG_ABBREVS:
            abbrev_counts[cand] = abbrev_counts.get(cand, 0) + 1
    new_abbrevs: set[str] = {a for a, cnt in abbrev_counts.items() if cnt >= 2}

    return {
        "Scientific Names":         sorted(new_sci),
        "Gene Names":               sorted(new_genes),
        "Protein Names":            sorted(new_proteins),
        "Event Names":              sorted(new_events),
        "Regulatory Abbreviations": sorted(new_abbrevs),
    }


def export_newterms_txt(
    terms: dict[str, list[str]],
    output_path: str,
    source_name: str = "",
) -> int:
    """
    Write _newterms.txt report.  Returns total number of candidate terms.
    """
    total = sum(len(v) for v in terms.values())
    with open(output_path, "w", encoding="utf-8") as fh:
        header = "# New term candidates"
        if source_name:
            header += f" — {source_name}"
        fh.write(header + "\n")
        fh.write(
            "# Review each section.  Confirmed terms should be added to the\n"
            "# known lists in clean_docx.py, and to latinTerms / specificFixes\n"
            "# in CleanUp.jsx / TitleCaseHeadings.jsx as appropriate.\n"
        )
        for category, term_list in terms.items():
            fh.write(f"\n=== {category} ===\n")
            if term_list:
                for t in term_list:
                    fh.write(f"  {t}\n")
            else:
                fh.write("  (none found)\n")
    return total


def main(argv: list[str]) -> None:
    if len(argv) < 2:
        print("Usage: python clean_docx.py INPUT_DOCX")
        sys.exit(1)

    source_path = argv[1]
    base, ext = os.path.splitext(source_path)
    clean_path         = base + "_clean"     + ext
    footnotes_path     = base + "_footnotes" + ext
    footnotes_txt_path = base + "_footnotes" + ".txt"
    newterms_path      = base + "_newterms"  + ".txt"

    # Footnotes extracted from original (unmodified) source
    try:
        n = extract_footnotes(source_path, footnotes_path)
        if n:
            print(f"Footnotes docx: {footnotes_path} ({n} notes)")
        else:
            print("Footnotes docx: none found")
    except Exception as e:
        print(f"Warning: could not extract footnotes docx: {e}", file=sys.stderr)

    try:
        n_txt = export_footnotes_txt(source_path, footnotes_txt_path)
        if n_txt:
            print(f"Footnotes txt:  {footnotes_txt_path} ({n_txt} notes)")
    except Exception as e:
        print(f"Warning: could not export footnotes txt: {e}", file=sys.stderr)

    # Scan for new terms (runs on original source before cleaning, to keep italic info)
    try:
        terms   = scan_new_terms(source_path)
        n_terms = export_newterms_txt(
            terms, newterms_path, source_name=os.path.basename(source_path)
        )
        n_cats = sum(1 for v in terms.values() if v)
        print(f"New terms:      {newterms_path} ({n_terms} candidates, {n_cats} categories)")
    except Exception as e:
        print(f"Warning: could not scan new terms: {e}", file=sys.stderr)

    # Apply all pre-processing and save with tables intact
    try:
        doc = Document(source_path)
        _strip_field_codes(doc)
        _mark_superscripts(doc)           # must run before _strip_character_styles strips rStyle
        _infer_paragraph_styles(doc)      # must run before _strip_character_styles strips bold/small-caps
        _strip_character_styles(doc)
        _clean_runs(doc)
        _extract_description_rows(doc)
        _register_assigned_styles(doc)    # must run before save so styles.xml includes our names
        doc.save(clean_path)
        print(f"Clean file: {clean_path}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main(sys.argv)
