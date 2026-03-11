from __future__ import annotations

import sys
import os
import copy
import re
import tempfile

# Ensure we can import locally installed packages (in .python-packages)
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOCAL_PACKAGES = os.path.join(SCRIPT_DIR, ".python-packages")
if os.path.isdir(LOCAL_PACKAGES) and LOCAL_PACKAGES not in sys.path:
    sys.path.insert(0, LOCAL_PACKAGES)

from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from lxml import etree  # type: ignore


def _get_description_row(table) -> str | None:
    """Return description text if first row is a merged title row, else None."""
    if not table.rows:
        return None
    first_row_texts = [(cell.text or "").strip() for cell in table.rows[0].cells]
    non_empty = [t for t in first_row_texts if t]
    if non_empty and len(set(non_empty)) == 1:
        return non_empty[0]
    return None


def _make_paragraph_element(text: str):
    """Create a bare <w:p> XML element containing the given plain text."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _mark_superscripts(doc: Document) -> None:
    """
    Walk every run in the document (body paragraphs and all table cells).
    - Superscript runs: wrap text as {{text}} and remove the superscript property.
    - Footnote references: replace the <w:footnoteReference> element with {{fn:N}}.
    - Endnote references: replace the <w:endnoteReference> element with {{en:N}}.
    """
    fn_tag = qn("w:footnoteReference")
    en_tag = qn("w:endnoteReference")
    va_tag = qn("w:vertAlign")
    rpr_tag = qn("w:rPr")

    def process_run(run):
        r_el = run._r

        # --- Footnote reference ---
        fn_ref = r_el.find(fn_tag)
        if fn_ref is not None:
            fn_id = fn_ref.get(qn("w:id"), "?")
            r_el.remove(fn_ref)
            # Strip superscript from rPr so it renders as plain text
            rpr = r_el.find(rpr_tag)
            if rpr is not None:
                va = rpr.find(va_tag)
                if va is not None:
                    rpr.remove(va)
            t = OxmlElement("w:t")
            t.text = f"{{{{fn:{fn_id}}}}}"
            r_el.append(t)
            return

        # --- Endnote reference ---
        en_ref = r_el.find(en_tag)
        if en_ref is not None:
            en_id = en_ref.get(qn("w:id"), "?")
            r_el.remove(en_ref)
            rpr = r_el.find(rpr_tag)
            if rpr is not None:
                va = rpr.find(va_tag)
                if va is not None:
                    rpr.remove(va)
            t = OxmlElement("w:t")
            t.text = f"{{{{en:{en_id}}}}}"
            r_el.append(t)
            return

        # --- Regular superscript run ---
        if run.font.superscript:
            run.text = f"{{{{{run.text}}}}}"
            run.font.superscript = False

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                process_run(run)

    # Body paragraphs
    process_paragraphs(doc.paragraphs)

    # All table cells (doc.tables includes nested tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)


def _clean_runs(doc: Document) -> None:
    """
    Text-level cleanup applied to every run in the document before extraction.
    - Strip leading bullet character (U+2022 + space) from paragraph starts.
    - Replace tilde operator (U+223C ∼) with standard tilde (~).
    """
    BULLET = "\u2022"
    TILDE_OP = "\u223c"

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            # Strip leading bullet from the paragraph's first non-empty run
            for run in para.runs:
                if run.text.startswith(BULLET + " "):
                    run.text = run.text[2:]
                elif run.text.startswith(BULLET):
                    run.text = run.text[1:]
                break  # only check the first run for bullets

            # Replace tilde operator in all runs
            for run in para.runs:
                if TILDE_OP in run.text:
                    run.text = run.text.replace(TILDE_OP, "~")

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
                # Normalize lone "X" approval markers to lowercase
                if cell.text.strip() == "X":
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace("X", "x")
                # Strip direct bold formatting so InDesign paragraph style controls weight
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = None


_FOOTNOTES_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
_ENDNOTES_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
)
# w:type values that are Word internals, not user footnotes
_SKIP_FOOTNOTE_TYPES = {"separator", "continuationSeparator", "continuationNotice"}


def extract_footnotes(source_path: str, output_path: str) -> int:
    """
    Extract all footnotes (and endnotes) from source_path into a new .docx.
    Each entry is labelled with its {{fn:N}} / {{en:N}} marker so it matches
    the placeholders in the _text and _tables files.
    Returns the total number of notes written (0 = nothing to write, no file saved).
    """
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    doc = Document(source_path)
    out_doc = Document()
    body = out_doc.element.body
    count = 0

    # Remove the default empty paragraph before we start appending content
    # (doing it at the end risks removing our first label instead)
    default_p = body.find(qn("w:p"))
    if default_p is not None:
        body.remove(default_p)

    fn_tag    = qn("w:footnote")
    en_tag    = qn("w:endnote")
    id_attr   = qn("w:id")
    type_attr = qn("w:type")
    t_tag     = qn("w:t")
    # <w:footnoteRef/> and <w:endnoteRef/> auto-number markers — strip from output
    ref_tags = {qn("w:footnoteRef"), qn("w:endnoteRef")}

    rpr_tag   = qn("w:rPr")
    keep_rpr  = {qn("w:i"), qn("w:iCs")}  # formatting to preserve

    def _copy_note_paragraphs(note_el):
        """Deep-copy a note's <w:p> children, stripping the auto-number run,
        leading whitespace, and all run formatting except italics."""
        for p_el in note_el.findall(qn("w:p")):
            p_copy = copy.deepcopy(p_el)
            # Remove any run that contains only a footnoteRef/endnoteRef element
            for r_el in p_copy.findall(qn("w:r")):
                children = list(r_el)
                non_rpr = [c for c in children if c.tag != rpr_tag]
                if len(non_rpr) == 1 and non_rpr[0].tag in ref_tags:
                    r_el.getparent().remove(r_el)
            # Strip leading whitespace from the first remaining run (Word inserts a
            # tab/space between the footnoteRef marker and the actual note text)
            remaining = p_copy.findall(qn("w:r"))
            if remaining:
                t_el = remaining[0].find(t_tag)
                if t_el is not None and t_el.text:
                    stripped = t_el.text.lstrip(" \t\xa0")
                    if stripped:
                        t_el.text = stripped
                    else:
                        remaining[0].getparent().remove(remaining[0])
            # Strip all run formatting except italics
            for r_el in p_copy.findall(qn("w:r")):
                rpr = r_el.find(rpr_tag)
                if rpr is not None:
                    for child in list(rpr):
                        if child.tag not in keep_rpr:
                            rpr.remove(child)
                    if len(rpr) == 0:
                        r_el.remove(rpr)
            body.append(p_copy)

    def _extract_from_part(rel_uri, note_tag, label_prefix):
        nonlocal count
        try:
            part = doc.part.part_related_by(rel_uri)
            part_xml = etree.fromstring(part.blob)
        except (KeyError, Exception):
            return  # this document has no footnotes/endnotes
        for note in part_xml.findall(note_tag):
            if note.get(type_attr) in _SKIP_FOOTNOTE_TYPES:
                continue
            note_id = note.get(id_attr, "?")
            body.append(_make_paragraph_element(f"{{{{{label_prefix}:{note_id}}}}}"))
            _copy_note_paragraphs(note)
            body.append(_make_paragraph_element(""))
            count += 1

    _extract_from_part(_FOOTNOTES_REL, fn_tag, "fn")
    _extract_from_part(_ENDNOTES_REL,  en_tag, "en")

    if count == 0:
        return 0

    out_doc.save(output_path)
    return count


def extract_tables(source_path: str, output_path: str) -> None:
    """
    Read a source .docx file and write a new .docx containing only the tables
    (in order). Merged cells and formatting are preserved via XML copy.
    """
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    src_doc = Document(source_path)
    out_doc = Document()
    body = out_doc.element.body

    for table_index, table in enumerate(src_doc.tables, start=1):
        body.append(_make_paragraph_element(f"Table {table_index}"))

        description = _get_description_row(table)
        if description is not None:
            body.append(_make_paragraph_element(description))

        # Deep-copy the table XML to preserve merges and formatting
        tbl_copy = copy.deepcopy(table._tbl)

        # Remove the description row from the copy if it was pulled out above
        if description is not None:
            first_tr = tbl_copy.find(qn("w:tr"))
            if first_tr is not None:
                tbl_copy.remove(first_tr)

        body.append(tbl_copy)
        body.append(_make_paragraph_element(""))

    # Remove the default empty paragraph Word adds to new documents
    default_p = body.find(qn("w:p"))
    if default_p is not None:
        body.remove(default_p)

    out_doc.save(output_path)


def extract_text(source_path: str, output_path: str) -> None:
    """
    Read a source .docx file and write a copy where every table is replaced
    by its label ("Table N") and, if present, the description row text.
    All other content is preserved as-is.
    """
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    doc = Document(source_path)
    body = doc.element.body

    tbl_tag = qn("w:tbl")
    table_objects = list(doc.tables)
    tbl_elements = [el for el in body if el.tag == tbl_tag]

    for table_index, (tbl_el, table) in enumerate(zip(tbl_elements, table_objects), start=1):
        label_el = _make_paragraph_element(f"Table {table_index}")
        tbl_el.addprevious(label_el)

        description = _get_description_row(table)
        if description is not None:
            description = re.sub(r"^(Table\s+\d+)\.\s+", r"\1: ", description)
            desc_el = _make_paragraph_element(description)
            label_el.addnext(desc_el)

        tbl_el.getparent().remove(tbl_el)

    doc.save(output_path)


def main(argv: list[str]) -> None:
    if len(argv) < 2:
        print("Usage: python extract_tables.py INPUT_DOCX")
        sys.exit(1)

    source_path = argv[1]
    base, ext = os.path.splitext(source_path)
    tables_path    = base + "_tables"    + ext
    text_path      = base + "_text"      + ext
    footnotes_path = base + "_footnotes" + ext

    # Footnotes are extracted from the original source (unmodified)
    try:
        n = extract_footnotes(source_path, footnotes_path)
        if n:
            print(f"Footnotes file: {footnotes_path} ({n} notes)")
        else:
            print("Footnotes file: none found")
    except Exception as e:
        print(f"Warning: could not extract footnotes: {e}", file=sys.stderr)

    # Pre-process: mark superscripts/footnotes in a temp copy, then extract from that
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=ext)
    os.close(tmp_fd)
    try:
        doc = Document(source_path)
        _mark_superscripts(doc)
        _clean_runs(doc)
        doc.save(tmp_path)

        extract_tables(tmp_path, tables_path)
        print(f"Tables file: {tables_path}")
        extract_text(tmp_path, text_path)
        print(f"Text file:   {text_path}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    finally:
        os.unlink(tmp_path)


if __name__ == "__main__":
    main(sys.argv)
